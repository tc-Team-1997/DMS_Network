"""Page-level OCR via Tesseract (+ pdf2image for PDFs), with an optional
vision-language fallback (Qwen2.5-VL / Llava) for low-confidence scans.

Output is structured (page index, text, confidence, word-level bboxes optional)
so downstream consumers (classifier, NER, layout reasoning) can trace any
inference back to a specific region of the source document.

Office-format text extraction (DOCX / XLSX) is handled here so that all
extract-text-from-upload logic lives in one place.
"""
from __future__ import annotations

import io
import logging
import os
from dataclasses import dataclass, field
from typing import List

import pytesseract
from PIL import Image

log = logging.getLogger(__name__)


@dataclass
class OcrPage:
    page: int
    text: str
    mean_confidence: float   # 0..100
    width: int               # px
    height: int              # px


@dataclass
class OcrResult:
    pages:         List[OcrPage]
    full_text:     str
    languages:     List[str]
    mean_confidence: float
    backend:       str = "tesseract"   # "tesseract" | "<vision-model>" | "passthrough"


def _lang_config() -> str:
    """
    Tesseract language pack selector. Always include English + Arabic in
    banking-KYC contexts (NBE, Gulf). Extend per tenant config later.
    """
    return os.environ.get("OCR_LANGS", "eng")


def _detect_languages(text: str) -> List[str]:
    """Lightweight language tag — OCR output, not a full langdet."""
    if any("\u0600" <= c <= "\u06FF" for c in text):
        return ["ara", "eng"]
    return ["eng"]


def _ocr_image(data: bytes, page: int = 1) -> OcrPage:
    img = Image.open(io.BytesIO(data))
    # pytesseract `image_to_data` gives us confidence per word; avg them.
    try:
        data_dict = pytesseract.image_to_data(
            img, lang=_lang_config(), output_type=pytesseract.Output.DICT,
        )
    except pytesseract.TesseractNotFoundError:
        log.error("tesseract binary not found; ensure TESSERACT_CMD is set")
        raise
    confs = [int(c) for c in data_dict["conf"] if c != "-1"]
    mean_conf = (sum(confs) / len(confs)) if confs else 0.0
    text = pytesseract.image_to_string(img, lang=_lang_config()).strip()
    return OcrPage(
        page=page, text=text, mean_confidence=float(mean_conf),
        width=img.width, height=img.height,
    )


def _ocr_pdf(data: bytes) -> List[OcrPage]:
    """Raster every page, OCR each. Rasterisation dependency lives here only."""
    from pdf2image import convert_from_bytes
    images = convert_from_bytes(
        data,
        dpi=200,
        poppler_path=os.environ.get("POPPLER_PATH") or None,
    )
    out: List[OcrPage] = []
    for i, img in enumerate(images, start=1):
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        out.append(_ocr_image(buf.getvalue(), page=i))
    return out


_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file using python-docx."""
    try:
        from docx import Document as DocxDocument  # type: ignore[import]
    except ImportError:
        log.warning("python-docx not installed; returning empty string for DOCX")
        return ""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text]
        # Also extract text from tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text not in paragraphs:
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as exc:  # noqa: BLE001
        log.warning("extract_docx_text failed: %s", exc)
        return ""


def extract_xlsx_text(file_bytes: bytes) -> str:
    """Extract plain text from a .xlsx file using openpyxl."""
    try:
        import openpyxl  # type: ignore[import]
    except ImportError:
        log.warning("openpyxl not installed; returning empty string for XLSX")
        return ""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        parts: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) for cell in row if cell is not None)
                if row_text.strip():
                    parts.append(row_text)
        wb.close()
        return "\n".join(parts).strip()
    except Exception as exc:  # noqa: BLE001
        log.warning("extract_xlsx_text failed: %s", exc)
        return ""


def _tesseract_ocr(data: bytes, mime_type: str) -> OcrResult:
    """The pure-Tesseract path, isolated so the orchestrator can fall back."""
    if mime_type == "application/pdf":
        pages = _ocr_pdf(data)
    elif mime_type.startswith("image/"):
        pages = [_ocr_image(data)]
    elif mime_type == _DOCX_MIME or mime_type in (
        "application/msword",
        "application/vnd.ms-word",
    ):
        text = extract_docx_text(data)
        return OcrResult(
            pages=[OcrPage(page=1, text=text, mean_confidence=100.0, width=0, height=0)],
            full_text=text,
            languages=_detect_languages(text),
            mean_confidence=100.0,
            backend="docx",
        )
    elif mime_type == _XLSX_MIME or mime_type in (
        "application/vnd.ms-excel",
        "application/octet-stream",
    ):
        # application/octet-stream is a common fallback for xlsx uploads; we
        # try openpyxl and fall back gracefully if the bytes aren't a workbook.
        if mime_type == "application/octet-stream":
            text = ""
            try:
                text = extract_xlsx_text(data)
            except Exception:  # noqa: BLE001
                pass
            if not text:
                # Not an xlsx — fall through to utf-8 passthrough below.
                try:
                    text = data.decode("utf-8", errors="replace")
                except Exception:  # noqa: BLE001
                    text = ""
        else:
            text = extract_xlsx_text(data)
        return OcrResult(
            pages=[OcrPage(page=1, text=text, mean_confidence=100.0, width=0, height=0)],
            full_text=text,
            languages=_detect_languages(text),
            mean_confidence=100.0,
            backend="xlsx",
        )
    else:
        # Plain text — utf-8 passthrough.
        try:
            text = data.decode("utf-8", errors="replace")
        except Exception:  # noqa: BLE001
            text = ""
        return OcrResult(
            pages=[OcrPage(page=1, text=text, mean_confidence=100.0, width=0, height=0)],
            full_text=text,
            languages=_detect_languages(text),
            mean_confidence=100.0,
            backend="passthrough",
        )

    full_text = "\n\n".join(p.text for p in pages if p.text).strip()
    mean_conf = (sum(p.mean_confidence for p in pages) / len(pages)) if pages else 0.0
    return OcrResult(
        pages=pages,
        full_text=full_text,
        languages=_detect_languages(full_text),
        mean_confidence=round(mean_conf, 2),
        backend="tesseract",
    )


def ocr_document(data: bytes, mime_type: str) -> OcrResult:
    """Top-level entry. Tesseract first; if confidence is poor AND a vision
    model is configured, re-run via Qwen-VL / Llava and take the better
    output. Everything downstream (classify → extract → embed) is oblivious
    to which backend produced the text."""
    if not data:
        return OcrResult(pages=[], full_text="", languages=[], mean_confidence=0.0)

    tess = _tesseract_ocr(data, mime_type)
    if tess.backend == "passthrough":
        return tess

    # Lazy import so vision deps aren't required when the feature is off.
    try:
        from .vision import vision_available, vision_ocr_document, choose_best
    except Exception:  # noqa: BLE001
        return tess
    if not vision_available():
        return tess

    # Only bother with vision when Tesseract's output is suspicious.
    below_threshold = tess.mean_confidence < float(
        os.environ.get("DOCBRAIN_VISION_OCR_THRESHOLD", "70")
    )
    short_text = len((tess.full_text or "").strip()) < 120
    if not (below_threshold or short_text):
        return tess

    try:
        vis = vision_ocr_document(data, mime_type)
    except Exception as exc:  # noqa: BLE001
        log.warning("vision ocr failed, keeping tesseract: %s", exc)
        return tess

    best, backend = choose_best(tess, vis)
    # Tag the chosen result with the backend that produced it.
    return OcrResult(
        pages=best.pages,
        full_text=best.full_text,
        languages=best.languages,
        mean_confidence=best.mean_confidence,
        backend=backend,
    )
